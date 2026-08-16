import os
import json
import re
import io
import requests
from flask import Flask, render_template, request, jsonify, send_file
from dotenv import load_dotenv
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor

load_dotenv()

app = Flask(__name__)
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
print("KEY LOADED:", GEMINI_API_KEY[:10] if GEMINI_API_KEY else "NONE FOUND")

GEMINI_URL = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-flash-lite-latest:generateContent?key={GEMINI_API_KEY}"


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/form")
def form():
    return render_template("form.html")


@app.route("/api/analyze", methods=["POST"])
def analyze():
    data = request.get_json()

    company_name = data.get("companyName", "").strip()
    company_domain = data.get("companyDomain", "").strip()
    competitors = data.get("competitors", [])

    if not company_name or not competitors:
        return jsonify({"error": "Missing company name or competitors"}), 400

    comp_list = "; ".join(
        f"{c.get('name','')} ({c.get('domain','')})" if c.get("domain") else c.get("name", "")
        for c in competitors
    )

    prompt = f"""You are a competitive intelligence analyst. Use your knowledge to research the following company and its competitors, then respond with ONLY valid JSON — no markdown fences, no commentary before or after — matching exactly this schema:

{{
  "company": {{"name": string, "summary": string (max 20 words)}},
  "competitors": [
    {{
      "name": string,
      "domain": string,
      "threat_level": "Low" | "Medium" | "High",
      "positioning": string (max 22 words),
      "strengths": [string, string] (max 7 words each),
      "weaknesses": [string, string] (max 7 words each),
      "recent_moves": [string, string] (max 10 words each),
      "pricing_signal": string (max 14 words)
    }}
  ],
  "comparison_matrix": {{
    "dimensions": [string, string, string, string],
    "rows": {{ "CompanyOrCompetitorName": [string, string, string, string] (max 5 words each, one row per company AND per competitor) }}
  }},
"recommendations": [string, string, string] (max 16 words each),
  "strategic_advice": string (max 60 words, a consultant-style paragraph giving overall strategic direction)
}}
Keep every string short, punchy, analyst-briefing tone. No fluff.

Subject company: {company_name}{f' ({company_domain})' if company_domain else ''}.
Competitors to investigate: {comp_list}."""

    try:
        response = requests.post(
            GEMINI_URL,
            headers={"Content-Type": "application/json"},
            json={
                "contents": [{"parts": [{"text": prompt}]}],
            },
            timeout=90,
        )
        response.raise_for_status()
        result = response.json()

        candidates = result.get("candidates", [])
        if not candidates:
            return jsonify({"error": "No response from Gemini", "raw": result}), 502

        parts = candidates[0].get("content", {}).get("parts", [])
        full_text = "\n".join(p.get("text", "") for p in parts).strip()

        full_text = re.sub(r"^```json\s*", "", full_text)
        full_text = re.sub(r"^```\s*", "", full_text)
        full_text = re.sub(r"```\s*$", "", full_text)

        first_brace = full_text.find("{")
        last_brace = full_text.rfind("}")
        if first_brace == -1 or last_brace == -1:
            return jsonify({"error": "No JSON found in AI response", "raw_text": full_text}), 502

        json_str = full_text[first_brace:last_brace + 1]
        parsed = json.loads(json_str)
        return jsonify(parsed)

    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"Request to Gemini API failed: {str(e)}"}), 502
    except json.JSONDecodeError as e:
        return jsonify({"error": f"Could not parse AI response as JSON: {str(e)}", "raw_text": full_text}), 502


def build_report_lines(data):
    """Turns report JSON into a flat list of (type, text) lines for export."""
    lines = []
    company = data.get("company", {})
    lines.append(("h1", f"Case File — {company.get('name', '')}"))
    lines.append(("p", company.get("summary", "")))
    lines.append(("h2", "Competitor Profiles"))

    for c in data.get("competitors", []):
        lines.append(("h3", f"{c.get('name','')} ({c.get('domain','')}) — Threat: {c.get('threat_level','')}"))
        lines.append(("p", f"Positioning: {c.get('positioning','')}"))
        lines.append(("p", "Strengths: " + "; ".join(c.get("strengths", []))))
        lines.append(("p", "Weaknesses: " + "; ".join(c.get("weaknesses", []))))
        lines.append(("p", "Recent moves: " + "; ".join(c.get("recent_moves", []))))
        lines.append(("p", f"Pricing: {c.get('pricing_signal','')}"))

    recs = data.get("recommendations", [])
    if recs:
        lines.append(("h2", "Recommendations"))
        for i, r in enumerate(recs, 1):
            lines.append(("p", f"{i}. {r}"))

    advice = data.get("strategic_advice", "")
    if advice:
        lines.append(("h2", "Strategic Advice"))
        lines.append(("p", advice))

    return lines

@app.route("/api/export/pdf", methods=["POST"])
def export_pdf():
    data = request.get_json()
    lines = build_report_lines(data)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)

    for kind, text in lines:
        safe_text = text.encode("latin-1", "replace").decode("latin-1")
        if not safe_text.strip():
            continue

        pdf.set_x(pdf.l_margin)

        if kind == "h1":
            pdf.set_font("Helvetica", "B", 18)
            pdf.multi_cell(0, 10, safe_text, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
        elif kind == "h2":
            pdf.ln(4)
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "B", 14)
            pdf.multi_cell(0, 8, safe_text, new_x="LMARGIN", new_y="NEXT")
        elif kind == "h3":
            pdf.ln(3)
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 7, safe_text, new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, safe_text, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)

    pdf_bytes = bytes(pdf.output())
    buffer = io.BytesIO(pdf_bytes)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="case_file.pdf",
        mimetype="application/pdf",
    )
@app.route("/api/export/docx", methods=["POST"])
def export_docx():
    data = request.get_json()
    lines = build_report_lines(data)

    doc = Document()

    for kind, text in lines:
        if kind == "h1":
            h = doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=2)
        elif kind == "h3":
            doc.add_heading(text, level=3)
        else:
            doc.add_paragraph(text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="case_file.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=False, host="0.0.0.0", port=port)
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=False, host="0.0.0.0", port=port)
